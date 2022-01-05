from docx import Document
import re
import os
import subprocess

def normalizeFile(path):
    fileList = os.listdir(path)
    print(f"[INFO] files name before normalization are: {fileList}")
    for filename in fileList:
        tmp_name = filename
        #del all blank space
        filename=filename.replace(' ','_')
        os.rename(path+'/'+tmp_name,path+'/'+filename)
    fileList = os.listdir(path)
    print(f'[INFO] files name after normalization are: {fileList}')

def formatfile(path):
    fileList = os.listdir(path)
    for filename in fileList:
        f = path+'/'+filename
        subprocess.run(['soffice', '--invisible', '--convert-to', 'docx', f, '--outdir', path])
        os.remove(f)

def replaceText(document,content,tex):
    all_tables = document.tables
    for tables in all_tables:
        for row in tables.rows:
            #print(row)
            for icell,cell in enumerate(row.cells):
                #print(cell.text)
                if cell.text.find(f"{content}")!=-1:
                    #print(f"loc is {icell} objet is 13641788003")
                    cell.text = f'{tex}'

    all_paragraphs = document.paragraphs
    for paragraphs in all_paragraphs:
        if(paragraphs.text.find(f"{content}")!=-1):
            newtext = re.sub(f"{content}",f'{tex}',paragraphs.text)
            paragraphs.text = newtext

def do_replace(dir,content,text):
    normalizeFile(dir)
    formatfile(dir)
    fileList = os.listdir(dir)
    for file in fileList:
        document = Document(dir+'/'+file)
        replaceText(document,content,text)
        tepname = dir+'/'+file.split('.docx')[0]+'_mod.docx'
        document.save(tepname)
        print(f"===================[INFO] replace {content} by {text} with {file} successfully!!!===========================")
        #remove original file and rename to original name
        originalname = dir+'/'+file
        os.remove(originalname)
        os.rename(tepname,originalname)

def replace_filename(dir,original,fixed):
    fileList = os.listdir(dir)
    for file in fileList:
        f = file.replace(original,fixed)
        print(f"===================[INFO] change name {file} to {f} successfully!!!=================================")
        os.rename(dir+'/'+file,dir+'/'+f)

if __name__ == "__main__":
    print("====================Welcome To Word Operation Script.=============================================================")
    print("====================With this new version, we can convert .doc to .docx now!!! So you can submit .doc now!!!======")
    print("====================More functions are coming soon !!!============================================================")

    import argparse
    parser = argparse.ArgumentParser('word operation')
    parser.add_argument('--dir', '-d',default='word', help='Path to the base directory of docx store.')
    parser.add_argument('--content', '-c',default='', help='contents that you want to be replaced.')
    parser.add_argument('--text', '-t',default='t', help='New contents that you want')
    parser.add_argument('--replacefilename', '-r',default='not',choices=['do', 'not'], help='If you want to reaname your file name, please enter -r do')
    parser.add_argument('--originaltext', '-o',default='', help='enter the text that you want to be replaced in the file name')
    parser.add_argument('--fixedtext', '-f',default='', help='test that you want to change to with file name')
    args = parser.parse_args()

    do_replace(args.dir,args.content,args.text)

    if(args.replacefilename=='do'):
        replace_filename(args.dir,args.originaltext,args.fixedtext)
##example
##python word_operation.py -d Jiaxing_Manager_20220105 -c 208118663 -t 211406901 -r do -o 202112 -f 202201
